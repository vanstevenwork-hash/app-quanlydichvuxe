from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/sivan/Downloads/5240045_SITV_CNTT1_28.1_BaoCaoDATN.docx")
BACKUP_PATH = DOCX_PATH.with_suffix(".backup-before-autofix.docx")


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_default_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def format_paragraph(paragraph, *, align=None, first_line=0.75, after=6, before=0, line=1.5):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if align is not None:
        paragraph.alignment = align


def add_paragraph(document: Document, text: str = "", *, bold=False, italic=False,
                  size=14, align=None, style="Normal", first_line=0.75,
                  after=6, before=0, line=1.5):
    p = document.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    format_paragraph(p, align=align, first_line=first_line, after=after, before=before, line=line)
    return p


def add_bullet(document: Document, text: str):
    p = document.add_paragraph(style="List Paragraph")
    p.add_run(text)
    format_paragraph(p, first_line=0, after=3)
    return p


def add_heading(document: Document, text: str, level: int):
    p = document.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16 if level == 1 else 14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    return p


def add_page_break(document: Document):
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table_of_contents(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run = OxmlElement("w:r")
    sep_text = OxmlElement("w:t")
    sep_text.text = "Mục lục sẽ được Word cập nhật khi mở tài liệu."
    sep_run.append(sep_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(sep_run)
    run._r.append(fld_end)
    format_paragraph(p, first_line=0, after=6)


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, head in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = head
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(13)
                p.paragraph_format.line_spacing = 1.2
    return table


def add_cover(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    add_paragraph(document, "TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, after=3)
    add_paragraph(document, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, after=18)
    add_paragraph(document, "", first_line=0, after=36)
    add_paragraph(document, "ĐỒ ÁN TỐT NGHIỆP", bold=True, size=18,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, after=6)
    add_paragraph(document, "ĐỀ TÀI", bold=True, size=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, after=12)
    add_paragraph(document, "XÂY DỰNG WEBSITE ĐẶT LỊCH SỬA CHỮA Ô TÔ TRỰC TUYẾN", bold=True, size=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, after=30, line=1.2)

    info_rows = [
        ("Giảng viên hướng dẫn", ": TS. Bùi Ngọc Dũng"),
        ("Sinh viên thực hiện", ": Tạ Văn Sĩ"),
        ("Mã sinh viên", ": 5240045"),
        ("Lớp", ": Công nghệ thông tin 1"),
        ("Khóa", ": 28.1"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.autofit = True
    for left, right in info_rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(14)
                    if idx == 0:
                        r.bold = True
    add_paragraph(document, "", first_line=0, after=30)
    add_paragraph(document, "Hà Nội - 2026", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, before=18)


def build_document(document: Document):
    add_cover(document)
    add_page_break(document)

    add_heading(document, "LỜI CẢM ƠN", 1)
    add_paragraph(
        document,
        "Trong quá trình thực hiện đồ án tốt nghiệp với đề tài xây dựng Website đặt lịch sửa chữa ô tô trực tuyến, em đã nhận được rất nhiều sự quan tâm, hướng dẫn và hỗ trợ quý báu từ thầy cô, gia đình và bạn bè. Đây là nguồn động lực quan trọng giúp em hoàn thành đề tài đúng định hướng và tiến độ.",
    )
    add_paragraph(
        document,
        "Em xin bày tỏ lòng biết ơn sâu sắc tới TS. Bùi Ngọc Dũng, người đã trực tiếp hướng dẫn, góp ý chuyên môn và định hướng cho em trong suốt quá trình phân tích, thiết kế, triển khai và hoàn thiện hệ thống. Những góp ý của thầy đã giúp em hiểu rõ hơn về cách xây dựng một sản phẩm phần mềm theo hướng bài bản, gắn với nhu cầu thực tế.",
    )
    add_paragraph(
        document,
        "Em cũng xin chân thành cảm ơn các thầy cô trong Khoa Công nghệ thông tin, Trường Đại học Giao thông Vận tải đã trang bị cho em nền tảng kiến thức quan trọng trong suốt quá trình học tập. Bên cạnh đó, em xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em cả về tinh thần lẫn thời gian để hoàn thành đồ án này.",
    )
    add_paragraph(
        document,
        "Mặc dù đã rất cố gắng, do thời gian thực hiện và kinh nghiệm thực tế còn hạn chế nên đồ án khó tránh khỏi những thiếu sót. Em rất mong tiếp tục nhận được những ý kiến đóng góp quý báu để có thể hoàn thiện đề tài tốt hơn trong thời gian tới.",
    )

    add_page_break(document)
    add_heading(document, "MỤC LỤC", 1)
    add_table_of_contents(document)

    add_page_break(document)
    add_heading(document, "DANH MỤC CÁC TỪ VIẾT TẮT", 1)
    add_table(document,
              ["STT", "Từ viết tắt", "Ý nghĩa"],
              [
                  ["1", "API", "Application Programming Interface"],
                  ["2", "CSDL", "Cơ sở dữ liệu"],
                  ["3", "CRUD", "Create - Read - Update - Delete"],
                  ["4", "JWT", "JSON Web Token"],
                  ["5", "UI", "User Interface"],
                  ["6", "UX", "User Experience"],
                  ["7", "MVC", "Model - View - Controller"],
                  ["8", "KPI", "Key Performance Indicator"],
              ])

    add_page_break(document)
    add_heading(document, "MỞ ĐẦU", 1)
    add_paragraph(
        document,
        "Trong bối cảnh nhu cầu sử dụng ô tô ngày càng tăng, các gara và trung tâm dịch vụ phải tiếp nhận số lượng lớn yêu cầu bảo dưỡng, kiểm tra và sửa chữa mỗi ngày. Tuy nhiên, ở nhiều cơ sở, việc đăng ký lịch hẹn vẫn được thực hiện chủ yếu qua điện thoại, tin nhắn hoặc ghi chép thủ công. Cách làm này dễ dẫn đến trùng lịch, thiếu thông tin xe, khó phân công kỹ thuật viên và khó theo dõi tiến độ xử lý.",
    )
    add_paragraph(
        document,
        "Xuất phát từ nhu cầu đó, đề tài tập trung xây dựng một Website đặt lịch sửa chữa ô tô trực tuyến theo mô hình client - server. Hệ thống cho phép khách hàng chủ động chọn dịch vụ, chọn ngày giờ, khai báo thông tin phương tiện và gửi yêu cầu đặt lịch trực tuyến; đồng thời hỗ trợ kỹ thuật viên và quản trị viên theo dõi, xử lý và thống kê toàn bộ hoạt động sửa chữa trên cùng một nền tảng.",
    )
    add_paragraph(
        document,
        "Project hiện tại được phát triển theo hướng full-stack với Front-end sử dụng ReactJS và Back-end sử dụng NodeJS, ExpressJS, MongoDB. Ngoài các chức năng lõi như đăng nhập, quản lý dịch vụ, đặt lịch và cập nhật trạng thái, hệ thống còn tích hợp xác thực JWT, gửi email thông báo lịch hẹn bằng Nodemailer và thống kê trực quan bằng ChartJS.",
    )
    add_paragraph(
        document,
        "Nội dung đồ án được trình bày thành bốn chương chính. Chương 1 khảo sát nghiệp vụ và nhu cầu thực tế của bài toán; Chương 2 phân tích thiết kế hệ thống; Chương 3 trình bày quá trình phát triển website; Chương 4 trình bày các giải pháp bảo mật, kiểm thử và đánh giá kết quả đạt được.",
    )

    add_heading(document, "CHƯƠNG 1: KHẢO SÁT NGHIỆP VỤ BÀI TOÁN", 1)
    add_heading(document, "1.1. Bối cảnh và nhu cầu thực tế", 2)
    add_paragraph(
        document,
        "Dịch vụ sửa chữa và bảo dưỡng ô tô là lĩnh vực yêu cầu quản lý thông tin chính xác theo từng xe, từng lịch hẹn và từng nhân sự kỹ thuật. Trong thực tế, nhiều gara vẫn tiếp nhận khách bằng hình thức thủ công nên việc tổng hợp thông tin thường chậm, khó tra cứu và thiếu đồng bộ giữa bộ phận tiếp nhận với kỹ thuật viên.",
    )
    add_paragraph(
        document,
        "Khách hàng ngày nay có xu hướng muốn đặt lịch trước để chủ động thời gian, biết trước chi phí dịch vụ và hạn chế thời gian chờ tại gara. Điều này đặt ra nhu cầu cần có một hệ thống trực tuyến giúp khách hàng đăng ký nhanh, đồng thời giúp gara chuẩn bị nhân lực, phụ tùng và khung giờ phù hợp.",
    )

    add_heading(document, "1.2. Quy trình nghiệp vụ hiện tại và những hạn chế", 2)
    add_paragraph(
        document,
        "Quy trình truyền thống thường bao gồm các bước: khách hàng liên hệ gara, nhân viên ghi nhận yêu cầu, sắp xếp khung giờ, tiếp nhận xe, phân công kỹ thuật viên và theo dõi tiến độ xử lý. Mô hình này phụ thuộc nhiều vào thao tác thủ công và kinh nghiệm cá nhân của người quản lý.",
    )
    add_bullet(document, "Thông tin lịch hẹn dễ bị thiếu hoặc không đồng nhất giữa khách hàng và gara.")
    add_bullet(document, "Khó kiểm soát tình trạng trùng lịch hoặc quá tải kỹ thuật viên vào một số khung giờ.")
    add_bullet(document, "Khách hàng khó theo dõi lịch sử sửa chữa và trạng thái xử lý sau khi đã gửi yêu cầu.")
    add_bullet(document, "Quản trị viên khó tổng hợp số liệu theo tháng, theo trạng thái hoặc theo số lượng dịch vụ.")
    add_bullet(document, "Việc thông báo xác nhận hoặc thay đổi lịch hẹn chủ yếu làm thủ công, tốn thời gian.")

    add_heading(document, "1.3. Yêu cầu chức năng của hệ thống", 2)
    add_heading(document, "1.3.1. Phân hệ khách hàng", 3)
    add_bullet(document, "Đăng ký tài khoản, đăng nhập, xem và cập nhật thông tin cá nhân.")
    add_bullet(document, "Xem danh sách dịch vụ sửa chữa, bảo dưỡng, kiểm tra và thay thế.")
    add_bullet(document, "Đặt lịch theo ngày và giờ, khai báo hãng xe, dòng xe, năm sản xuất, biển số và ghi chú tình trạng xe.")
    add_bullet(document, "Theo dõi danh sách lịch hẹn, xem trạng thái xử lý và hủy lịch khi điều kiện cho phép.")

    add_heading(document, "1.3.2. Phân hệ kỹ thuật viên", 3)
    add_bullet(document, "Xem danh sách lịch hẹn được phân công.")
    add_bullet(document, "Cập nhật trạng thái thực hiện từ đã xác nhận sang đang sửa và hoàn thành.")
    add_bullet(document, "Theo dõi thông tin phương tiện, ghi chú và thời gian bắt đầu, hoàn thành thực tế.")

    add_heading(document, "1.3.3. Phân hệ quản trị viên", 3)
    add_bullet(document, "Quản lý dịch vụ theo nghiệp vụ CRUD.")
    add_bullet(document, "Quản lý lịch hẹn, phân công kỹ thuật viên và cập nhật trạng thái tổng thể.")
    add_bullet(document, "Quản lý khách hàng, tạo tài khoản kỹ thuật viên, khóa hoặc kích hoạt tài khoản.")
    add_bullet(document, "Quản lý kho phụ tùng thông qua module inventory.")
    add_bullet(document, "Xem báo cáo thống kê tổng quan bằng dashboard và biểu đồ.")

    add_heading(document, "1.4. Yêu cầu phi chức năng", 2)
    add_paragraph(
        document,
        "Bên cạnh yêu cầu nghiệp vụ, hệ thống cần đảm bảo giao diện thân thiện, dễ sử dụng trên cả máy tính và thiết bị di động; phản hồi đủ nhanh với các thao tác xem dịch vụ, đặt lịch và lọc dữ liệu; dữ liệu người dùng phải được bảo vệ; kiến trúc cần dễ mở rộng thêm dịch vụ, kỹ thuật viên hoặc module quản lý kho trong tương lai.",
    )

    add_heading(document, "1.5. Kết luận chương", 2)
    add_paragraph(
        document,
        "Qua khảo sát nghiệp vụ có thể thấy việc số hóa quy trình đặt lịch sửa chữa ô tô là cần thiết và có tính ứng dụng cao. Những yêu cầu thu thập được là cơ sở để tiến hành phân tích, thiết kế và xây dựng hệ thống ở các chương tiếp theo.",
    )

    add_heading(document, "CHƯƠNG 2: PHÂN TÍCH THIẾT KẾ HỆ THỐNG", 1)
    add_heading(document, "2.1. Kiến trúc tổng thể", 2)
    add_paragraph(
        document,
        "Hệ thống được xây dựng theo mô hình client - server. Phía client là ứng dụng web ReactJS chịu trách nhiệm hiển thị giao diện, điều hướng và tương tác người dùng. Phía server là ExpressJS triển khai các RESTful API xử lý đăng nhập, quản lý người dùng, dịch vụ, lịch hẹn, kho phụ tùng và thống kê. Dữ liệu được lưu trữ trên MongoDB dưới dạng các collection độc lập nhưng có liên kết thông qua ObjectId.",
    )
    add_paragraph(
        document,
        "Luồng xử lý chính được thiết kế như sau: người dùng thực hiện thao tác trên giao diện, Front-end gửi request qua Axios đến Back-end, Back-end kiểm tra xác thực JWT và phân quyền theo vai trò, sau đó thao tác với cơ sở dữ liệu và trả kết quả về client. Với các sự kiện như tạo lịch hẹn hoặc cập nhật trạng thái, server còn kích hoạt dịch vụ gửi email thông báo cho khách hàng.",
    )

    add_heading(document, "2.2. Tác nhân và ca sử dụng", 2)
    add_paragraph(
        document,
        "Ba tác nhân chính của hệ thống là Khách hàng, Kỹ thuật viên và Quản trị viên. Mỗi tác nhân được giới hạn phạm vi thao tác thông qua middleware xác thực và phân quyền trên server.",
    )
    add_bullet(document, "Khách hàng: đăng ký, đăng nhập, xem dịch vụ, đặt lịch, xem hồ sơ và lịch sử sửa chữa.")
    add_bullet(document, "Kỹ thuật viên: xem lịch được giao và cập nhật tiến độ sửa chữa.")
    add_bullet(document, "Quản trị viên: quản trị dữ liệu hệ thống, nhân sự, lịch hẹn, báo cáo và kho phụ tùng.")

    add_heading(document, "2.3. Thiết kế cơ sở dữ liệu", 2)
    add_paragraph(
        document,
        "Cơ sở dữ liệu MongoDB của project hiện tại gồm bốn thực thể quan trọng là User, Service, Appointment và Inventory. Thiết kế này phản ánh đầy đủ các nhu cầu quản trị tài khoản, dịch vụ, lịch hẹn sửa chữa và phụ tùng trong gara.",
    )
    add_heading(document, "2.3.1. Bảng người dùng (User)", 3)
    add_paragraph(
        document,
        "Thực thể User lưu các thông tin như họ tên, email, số điện thoại, mật khẩu, vai trò, địa chỉ, ảnh đại diện, chuyên môn và trạng thái hoạt động. Vai trò được giới hạn trong ba giá trị customer, technician và admin.",
    )
    add_heading(document, "2.3.2. Bảng dịch vụ (Service)", 3)
    add_paragraph(
        document,
        "Thực thể Service lưu tên dịch vụ, nhóm dịch vụ, mô tả, giá tiền, thời lượng thực hiện, hình ảnh và trạng thái hoạt động. Dữ liệu này được dùng ở cả phía khách hàng lẫn quản trị viên.",
    )
    add_heading(document, "2.3.3. Bảng lịch hẹn (Appointment)", 3)
    add_paragraph(
        document,
        "Appointment là thực thể trung tâm, chứa liên kết đến khách hàng, dịch vụ và kỹ thuật viên; đồng thời lưu ngày hẹn, giờ hẹn, thông tin phương tiện, ghi chú, tổng giá, trạng thái và các mốc thời gian thực tế. Trạng thái của lịch hẹn được chuẩn hóa thành pending, confirmed, in-progress, completed và cancelled.",
    )
    add_heading(document, "2.3.4. Bảng kho phụ tùng (Inventory)", 3)
    add_paragraph(
        document,
        "Inventory hỗ trợ quản lý phụ tùng trong kho thông qua các thuộc tính như tên, mã SKU, danh mục và thông tin tìm kiếm. Đây là phần mở rộng hợp lý giúp hệ thống tiến gần hơn tới một nền tảng quản trị gara hoàn chỉnh.",
    )

    add_heading(document, "2.4. Ràng buộc nghiệp vụ và bảo mật dữ liệu", 2)
    add_bullet(document, "Chỉ khách hàng đã đăng nhập mới được đặt lịch.")
    add_bullet(document, "Khách hàng chỉ xem và hủy được lịch của chính mình.")
    add_bullet(document, "Kỹ thuật viên chỉ được cập nhật trạng thái các lịch được phân công.")
    add_bullet(document, "Quản trị viên có toàn quyền quản lý dịch vụ, tài khoản, lịch hẹn, phụ tùng và báo cáo.")
    add_bullet(document, "Mật khẩu người dùng được mã hóa bằng bcryptjs trước khi lưu vào cơ sở dữ liệu.")
    add_bullet(document, "Tất cả các API cần bảo vệ đều đi qua middleware xác thực JWT.")

    add_heading(document, "2.5. Kết luận chương", 2)
    add_paragraph(
        document,
        "Phân tích thiết kế cho thấy kiến trúc và mô hình dữ liệu của hệ thống phù hợp với bài toán quản lý lịch sửa chữa ô tô trực tuyến. Thiết kế này tạo nền tảng để triển khai các chức năng ở mức mã nguồn một cách rõ ràng và dễ mở rộng.",
    )

    add_heading(document, "CHƯƠNG 3: PHÁT TRIỂN WEBSITE", 1)
    add_heading(document, "3.1. Công nghệ sử dụng", 2)
    add_paragraph(
        document,
        "Front-end được phát triển bằng ReactJS kết hợp React Router để điều hướng, Axios để gọi API, Bootstrap và các thành phần giao diện để tăng tốc triển khai. Back-end sử dụng NodeJS, ExpressJS và Mongoose để xây dựng REST API. MongoDB được dùng làm hệ quản trị cơ sở dữ liệu NoSQL. Hệ thống còn tích hợp JWT cho xác thực, Nodemailer cho gửi email và ChartJS cho thống kê trực quan.",
    )

    add_heading(document, "3.2. Triển khai phía giao diện người dùng", 2)
    add_paragraph(
        document,
        "Ứng dụng React được tổ chức theo layout cho khách hàng và quản trị viên. Các route chính gồm trang chủ, trang dịch vụ, liên hệ, đặt lịch, hồ sơ cá nhân, lịch kỹ thuật viên và khu vực quản trị. Cơ chế ProtectedRoute được sử dụng để kiểm soát quyền truy cập tương ứng với từng vai trò người dùng.",
    )
    add_paragraph(
        document,
        "Trang đặt lịch là một trong những điểm nhấn của hệ thống. Quy trình đặt lịch được chia thành bốn bước gồm chọn dịch vụ, chọn ngày giờ, nhập thông tin xe và ghi chú tình trạng xe. Giao diện có phần tóm tắt chi phí ước tính ở bên phải, giúp người dùng theo dõi rõ ràng trước khi xác nhận.",
    )
    add_paragraph(
        document,
        "Trang hồ sơ khách hàng cung cấp khả năng xem lịch sử sửa chữa, lọc theo trạng thái, theo dõi kỹ thuật viên phụ trách và hủy lịch trong những trường hợp phù hợp. Đối với kỹ thuật viên, màn hình lịch làm việc hiển thị danh sách xe được giao và cho phép cập nhật trạng thái nhanh chóng.",
    )

    add_heading(document, "3.3. Triển khai phía máy chủ", 2)
    add_paragraph(
        document,
        "Server Express được chia thành các tầng route, controller, model, middleware và service. Cách tổ chức này giúp phân tách rõ trách nhiệm giữa định tuyến, xử lý nghiệp vụ và thao tác dữ liệu. Các route chính bao gồm auth, users, services, appointments, inventory và stats.",
    )
    add_paragraph(
        document,
        "Khối auth xử lý đăng ký, đăng nhập, lấy thông tin người dùng hiện tại và đổi mật khẩu. Khối service đảm nhiệm tìm kiếm, phân trang và CRUD dịch vụ. Khối appointment xử lý tạo lịch hẹn, xem danh sách theo vai trò, xem chi tiết, cập nhật trạng thái, phân công kỹ thuật viên và hủy lịch. Khối stats tổng hợp chỉ số dashboard và dữ liệu biểu đồ theo tháng, theo trạng thái.",
    )

    add_heading(document, "3.4. Các chức năng nổi bật đã xây dựng", 2)
    add_bullet(document, "Đăng ký và đăng nhập có kiểm tra hợp lệ dữ liệu đầu vào.")
    add_bullet(document, "Đặt lịch với thông tin phương tiện đầy đủ: hãng xe, dòng xe, năm sản xuất và biển số.")
    add_bullet(document, "Phân công kỹ thuật viên cho từng lịch hẹn từ màn hình quản trị.")
    add_bullet(document, "Cập nhật trạng thái xử lý và ghi nhận thời gian bắt đầu, hoàn thành thực tế.")
    add_bullet(document, "Quản lý dịch vụ, khách hàng, kỹ thuật viên và kho phụ tùng.")
    add_bullet(document, "Dashboard thống kê tổng khách hàng, số lịch sửa chữa, số xe đang xử lý và doanh thu.")
    add_bullet(document, "Biểu đồ cột theo tháng và biểu đồ tròn theo trạng thái lịch hẹn.")
    add_bullet(document, "Gửi email xác nhận và cập nhật lịch hẹn tự động cho khách hàng.")

    add_heading(document, "3.5. Đánh giá kết quả phát triển", 2)
    add_paragraph(
        document,
        "Qua quá trình triển khai, hệ thống đã đáp ứng tương đối đầy đủ các yêu cầu đặt ra cho một website đặt lịch sửa chữa ô tô trực tuyến. Cấu trúc mã nguồn rõ ràng, các API bám sát nghiệp vụ và giao diện hỗ trợ tốt cho ba vai trò sử dụng. Việc bổ sung module inventory và phần dashboard giúp project vượt ra ngoài phạm vi đặt lịch đơn thuần, tiến gần tới một giải pháp quản trị gara thực tế hơn.",
    )

    add_heading(document, "3.6. Kết luận chương", 2)
    add_paragraph(
        document,
        "Chương này đã trình bày các thành phần công nghệ và cách thức xây dựng hệ thống ở cả phía client và server. Kết quả triển khai cho thấy project có tính khả thi, có thể tiếp tục mở rộng thêm thanh toán trực tuyến, quản lý phụ tùng nâng cao và lịch sử bảo dưỡng chuyên sâu trong tương lai.",
    )

    add_heading(document, "CHƯƠNG 4: BẢO MẬT VÀ KIỂM THỬ", 1)
    add_heading(document, "4.1. Giải pháp bảo mật", 2)
    add_paragraph(
        document,
        "Bảo mật là yêu cầu quan trọng đối với hệ thống quản lý lịch hẹn và thông tin khách hàng. Trong project, xác thực người dùng được thực hiện bằng JWT. Khi đăng nhập thành công, server phát sinh token chứa thông tin định danh người dùng; các request tiếp theo phải gửi token qua header Authorization để được truy cập tài nguyên bảo vệ.",
    )
    add_paragraph(
        document,
        "Ngoài ra, hệ thống sử dụng middleware authorize để kiểm soát quyền truy cập theo vai trò. Mật khẩu người dùng được băm bằng bcryptjs trước khi lưu vào cơ sở dữ liệu. Các dữ liệu đầu vào quan trọng tại khối đăng ký và đăng nhập được kiểm tra bằng express-validator nhằm hạn chế dữ liệu sai định dạng.",
    )

    add_heading(document, "4.2. Kiểm thử chức năng", 2)
    add_paragraph(
        document,
        "Hệ thống được kiểm tra theo hướng kiểm thử chức năng ở các luồng quan trọng. Các ca kiểm thử tập trung vào xác thực tài khoản, tìm kiếm dịch vụ, đặt lịch hẹn, phân công kỹ thuật viên, cập nhật trạng thái và thống kê báo cáo.",
    )
    add_table(
        document,
        ["Mã TC", "Mô tả kiểm thử", "Kết quả mong đợi"],
        [
            ["TC01", "Đăng ký tài khoản với dữ liệu hợp lệ", "Tạo tài khoản thành công và nhận token đăng nhập."],
            ["TC02", "Đăng nhập sai mật khẩu", "Hệ thống từ chối và trả về thông báo lỗi phù hợp."],
            ["TC03", "Khách hàng đặt lịch với đầy đủ thông tin", "Lịch hẹn được tạo ở trạng thái chờ xác nhận."],
            ["TC04", "Quản trị viên phân công kỹ thuật viên", "Lịch hẹn được gắn kỹ thuật viên và chuyển sang đã xác nhận."],
            ["TC05", "Kỹ thuật viên cập nhật trạng thái hoàn thành", "Hệ thống lưu thời gian hoàn thành và cập nhật dashboard."],
            ["TC06", "Khách hàng hủy lịch đã hoàn thành", "Hệ thống không cho phép hủy và trả về thông báo phù hợp."],
        ],
    )

    add_heading(document, "4.3. Đánh giá ưu điểm và hạn chế", 2)
    add_paragraph(
        document,
        "Ưu điểm của hệ thống là giao diện trực quan, kiến trúc tách lớp rõ ràng, nghiệp vụ phân vai cụ thể, thống kê được trực quan hóa và có tích hợp email thông báo tự động. Phần quản lý dịch vụ, lịch hẹn, nhân sự và kho phụ tùng tạo nên một nền tảng có tính ứng dụng thực tế.",
    )
    add_paragraph(
        document,
        "Bên cạnh đó, project vẫn còn một số hạn chế như chưa tích hợp thanh toán trực tuyến, chưa có cơ chế đặt lại mật khẩu qua email, chưa có kiểm thử tự động và chưa tích hợp phân bổ lịch thông minh theo tải công việc của kỹ thuật viên. Đây là những hướng phát triển phù hợp cho các phiên bản tiếp theo.",
    )

    add_heading(document, "4.4. Kết luận chương", 2)
    add_paragraph(
        document,
        "Các giải pháp bảo mật và kiểm thử hiện tại đáp ứng tốt phạm vi của đồ án tốt nghiệp. Hệ thống có nền tảng đủ ổn định để tiếp tục hoàn thiện và mở rộng theo nhu cầu triển khai thực tế tại các gara ô tô.",
    )

    add_heading(document, "KẾT LUẬN VÀ KIẾN NGHỊ", 1)
    add_paragraph(
        document,
        "Đề tài xây dựng Website đặt lịch sửa chữa ô tô trực tuyến đã hoàn thành được mục tiêu chính là số hóa quy trình tiếp nhận và quản lý lịch sửa chữa. Hệ thống hỗ trợ khách hàng đặt lịch thuận tiện hơn, hỗ trợ kỹ thuật viên theo dõi công việc và hỗ trợ quản trị viên quản lý tổng thể hoạt động của gara trên một nền tảng thống nhất.",
    )
    add_paragraph(
        document,
        "Từ kết quả của project hiện tại có thể khẳng định việc áp dụng các công nghệ web hiện đại như ReactJS, ExpressJS và MongoDB là phù hợp với bài toán. Cấu trúc dự án cũng cho phép mở rộng thêm nhiều chức năng như thanh toán trực tuyến, quản lý hóa đơn, nhắc lịch bảo dưỡng định kỳ, phân tích dữ liệu khách hàng và tối ưu lịch phân công kỹ thuật viên.",
    )
    add_paragraph(
        document,
        "Trong thời gian tới, nếu tiếp tục phát triển, hệ thống nên được hoàn thiện thêm về kiểm thử tự động, tối ưu trải nghiệm người dùng, tăng cường bảo mật ở mức triển khai thực tế và tích hợp sâu hơn với các quy trình vận hành của gara. Đây sẽ là tiền đề để sản phẩm có thể chuyển từ mô hình đồ án sang ứng dụng thực tế.",
    )

    add_heading(document, "DANH MỤC TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Tài liệu chính thức của ReactJS.",
        "Tài liệu chính thức của ExpressJS.",
        "Tài liệu chính thức của MongoDB và Mongoose.",
        "Tài liệu chính thức của NodeJS.",
        "Tài liệu chính thức của ChartJS.",
        "Tài liệu chính thức của Nodemailer.",
        "Mã nguồn project Website đặt lịch sửa chữa ô tô trực tuyến trong thư mục app-quanlydichvuxe.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_paragraph(document, f"[{idx}] {ref}", first_line=0, after=3)


def main():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    document = Document(DOCX_PATH)
    clear_document(document)
    set_default_font(document)
    build_document(document)
    document.save(DOCX_PATH)
    print(f"Updated: {DOCX_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
